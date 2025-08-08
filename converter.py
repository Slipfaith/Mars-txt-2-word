import logging
import os
from typing import Callable, Iterable, List, Optional

import chardet
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)


def detect_encoding(file_path: str, num_bytes: int = 10000) -> Optional[str]:
    """Detect file encoding using chardet with fallback for MacCyrillic."""
    with open(file_path, 'rb') as f:
        rawdata = f.read(num_bytes)
    result = chardet.detect(rawdata)
    encoding = result.get('encoding')
    if encoding and encoding.lower() == "maccyrillic":
        encoding = "cp1251"
    logger.debug("%s encoding detected as %s", file_path, encoding)
    return encoding


def read_lines_auto(
    file_path: str,
    default_encoding: str = 'utf-8',
    force_encoding: Optional[str] = None,
) -> List[str]:
    """Read lines from file using detected or forced encoding.

    Must-have: errors='replace', не убиваем пробелы — только \r\n.
    """
    encoding = force_encoding if force_encoding else detect_encoding(file_path) or default_encoding
    source = "forced" if force_encoding else "auto-detected"
    logger.debug("Reading %s with %s encoding %s", file_path, source, encoding)
    try:
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            # важно: не .strip(), чтобы не терять значащие пробелы/пустые строки
            lines = [line.rstrip("\r\n") for line in f]
        logger.debug("Read %d lines from %s", len(lines), file_path)
        return lines
    except Exception as exc:
        logger.error("Failed reading %s with %s: %s", file_path, encoding, exc)
        with open(file_path, 'r', encoding=default_encoding, errors='replace') as f:
            lines = [line.rstrip("\r\n") for line in f]
        logger.debug(
            "Fallback read %d lines from %s using %s", len(lines), file_path, default_encoding
        )
        return lines


def iter_block_items(parent) -> Iterable:
    """Yield Paragraph and Table items from a Word document."""
    parent_elm = parent.element.body if hasattr(parent, 'element') else parent
    for child in parent_elm:
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)


def _unique_path(path: str) -> str:
    """Защита от перезаписи: если путь занят — добавляем _1, _2, ..."""
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    i = 1
    while True:
        cand = f"{base}_{i}{ext}"
        if not os.path.exists(cand):
            return cand
        i += 1


def export_paths_to_word(
    eng_paths: List[str],
    rus_paths: List[str],
    output_path: str,
    rus_force_encoding: Optional[str] = None,
    progress: Optional[Callable[[int, int], None]] = None,
) -> None:
    """Create Word document from explicit file lists."""
    if not eng_paths or not rus_paths:
        raise ValueError("No files provided")

    for p in eng_paths + rus_paths:
        if not os.path.exists(p):
            raise FileNotFoundError(f"Path does not exist: {p}")

    eng_map = {os.path.basename(p): p for p in eng_paths}
    rus_map = {os.path.basename(p): p for p in rus_paths}

    # Проверка парности
    missing_rus = [name for name in eng_map.keys() if name not in rus_map]
    missing_eng = [name for name in rus_map.keys() if name not in eng_map]
    if missing_rus:
        logger.warning("Missing russian files for: %s", ", ".join(missing_rus))
    if missing_eng:
        logger.warning("Missing english files for: %s", ", ".join(missing_eng))

    # Определяем расширение по ENG списку (как у тебя)
    extensions = {os.path.splitext(p)[1].lower() for p in eng_paths}
    if len(extensions) != 1 or list(extensions)[0] not in ('.txt', '.srt'):
        raise ValueError("Ambiguous or unsupported file extensions in English files")
    file_extension = list(extensions)[0].lstrip('.')

    # Готовим DOCX
    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    row = table.add_row()
    row.cells[0].text = f"Тип файлов: {file_extension}"
    row.cells[1].text = ""

    paired = [fn for fn in eng_map.keys() if fn in rus_map]
    total = len(paired)
    for idx, filename in enumerate(paired, 1):
        eng_path = eng_map[filename]
        rus_path = rus_map[filename]

        eng_lines = read_lines_auto(eng_path, default_encoding="utf-8")
        rus_lines = read_lines_auto(
            rus_path,
            default_encoding="cp1251",
            force_encoding=rus_force_encoding,
        )

        # маркер файла
        marker_row = table.add_row()
        file_cell = marker_row.cells[0]
        file_cell.text = ""
        run = file_cell.paragraphs[0].add_run(f"Файл: {filename}")
        run.bold = True
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(r'<w:shd {} w:fill="CCFFCC"/>'.format(nsdecls("w")))
        file_cell._tc.get_or_add_tcPr().append(shading)
        marker_row.cells[1].text = ""

        # строки
        num_rows = max(len(eng_lines), len(rus_lines))
        for i in range(num_rows):
            data_row = table.add_row()
            data_row.cells[0].text = eng_lines[i] if i < len(eng_lines) else ""
            data_row.cells[1].text = rus_lines[i] if i < len(rus_lines) else ""

        if progress:
            progress(idx, total)

    doc.save(output_path)
    logger.info("Word document saved: %s", output_path)


def export_to_word(
    eng_folder: str,
    rus_folder: str,
    output_path: str,
    file_extension: Optional[str] = None,
    rus_force_encoding: Optional[str] = None,
    progress: Optional[Callable[[int, int], None]] = None,
) -> None:
    """Create Word document from pairs of subtitle files."""
    if not os.path.isdir(eng_folder) or not os.path.isdir(rus_folder):
        raise FileNotFoundError("Folders do not exist")

    files_in_eng = [f for f in os.listdir(eng_folder) if os.path.isfile(os.path.join(eng_folder, f))]
    files_in_rus = [f for f in os.listdir(rus_folder) if os.path.isfile(os.path.join(rus_folder, f))]

    if file_extension is None:
        # автоопределение по ENG
        exts = {os.path.splitext(f)[1].lower() for f in files_in_eng}
        exts = {e for e in exts if e in ('.txt', '.srt')}
        if len(exts) != 1:
            raise ValueError("Ambiguous or unsupported file extensions in English folder")
        file_extension = list(exts)[0].lstrip('.')

    eng_files = {f for f in files_in_eng if f.lower().endswith(f".{file_extension}")}
    rus_files = {f for f in files_in_rus if f.lower().endswith(f".{file_extension}")}

    # Проверка парности
    missing_rus = sorted(eng_files - rus_files)
    missing_eng = sorted(rus_files - eng_files)
    if missing_rus:
        logger.warning("Missing russian files for: %s", ", ".join(missing_rus))
    if missing_eng:
        logger.warning("Missing english files for: %s", ", ".join(missing_eng))

    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    row = table.add_row()
    row.cells[0].text = f"Тип файлов: {file_extension}"
    row.cells[1].text = ""

    paired = sorted(eng_files & rus_files)
    total = len(paired)
    for idx, filename in enumerate(paired, 1):
        eng_file_path = os.path.join(eng_folder, filename)
        rus_file_path = os.path.join(rus_folder, filename)
        if not os.path.exists(rus_file_path) or not os.path.exists(eng_file_path):
            logger.warning("Missing pair for %s", filename)
            continue
        eng_lines = read_lines_auto(eng_file_path, default_encoding='utf-8')
        rus_lines = read_lines_auto(
            rus_file_path, default_encoding='cp1251', force_encoding=rus_force_encoding
        )

        # row marking the beginning of a new file
        marker_row = table.add_row()
        file_cell = marker_row.cells[0]
        file_cell.text = ""
        run = file_cell.paragraphs[0].add_run(f"Файл: {filename}")
        run.bold = True
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(
            r'<w:shd {} w:fill="CCFFCC"/>'.format(nsdecls("w"))
        )
        file_cell._tc.get_or_add_tcPr().append(shading)
        marker_row.cells[1].text = ""

        num_rows = max(len(eng_lines), len(rus_lines))
        for i in range(num_rows):
            data_row = table.add_row()
            data_row.cells[0].text = eng_lines[i] if i < len(eng_lines) else ''
            data_row.cells[1].text = rus_lines[i] if i < len(rus_lines) else ''

        if progress:
            progress(idx, total)

    doc.save(output_path)
    logger.info("Word document saved: %s", output_path)


def import_from_word(
    word_path: str,
    eng_output_folder: str,
    rus_output_folder: str,
    overwrite: bool = False,
    progress: Optional[Callable[[int, int], None]] = None,
) -> None:
    """Разбить Word обратно на файлы.

    Must-have:
    - RU для .srt пишем в UTF-8 (фикс charmap), иначе cp1251.
    - Защита от перезаписи.
    - Прогресс.
    """
    if not os.path.isfile(word_path):
        raise FileNotFoundError(word_path)
    os.makedirs(eng_output_folder, exist_ok=True)
    os.makedirs(rus_output_folder, exist_ok=True)

    doc = Document(word_path)

    # определяем формат из первой строки таблицы
    file_format = "txt"
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.startswith("Тип файлов:"):
            file_format = table.rows[0].cells[0].text.split(":", 1)[1].strip().lower()
            break
    ext = "srt" if file_format == "srt" else "txt"
    rus_write_encoding = "utf-8" if ext == "srt" else "cp1251"

    # собираем данные
    file_data = {}
    current_filename: Optional[str] = None

    # твоя логика — проход по таблицам (в документе структура табличная)
    for table in doc.tables:
        for idx, row in enumerate(table.rows):
            first = row.cells[0].text.strip()
            second = row.cells[1].text.strip() if len(row.cells) > 1 else ""

            if idx == 0 and first.startswith("Тип файлов:"):
                # уже определили выше, просто пропустим
                continue

            if first.startswith("Файл:"):
                current_filename = first.replace("Файл:", "").strip()
                file_data[current_filename] = ([], [])
                logger.debug("Processing section for %s", current_filename)
                continue

            if current_filename:
                eng_lines, rus_lines = file_data[current_filename]
                eng_lines.append(first)
                rus_lines.append(second)

    # запись
    items = list(file_data.items())
    total = len(items)
    for i, (filename, (eng_lines, rus_lines)) in enumerate(items, 1):
        filename_with_ext = (
            filename if filename.lower().endswith(f".{ext}") else f"{filename}.{ext}"
        )
        eng_file_path = os.path.join(eng_output_folder, filename_with_ext)
        rus_file_path = os.path.join(rus_output_folder, filename_with_ext)

        if not overwrite:
            eng_file_path = _unique_path(eng_file_path)
            rus_file_path = _unique_path(rus_file_path)

        with open(eng_file_path, "w", encoding="utf-8", errors="replace") as f:
            f.write("\n".join(eng_lines))
        with open(rus_file_path, "w", encoding=rus_write_encoding, errors="replace") as f:
            f.write("\n".join(rus_lines))
        logger.debug("Saved %s and %s", eng_file_path, rus_file_path)

        if progress:
            progress(i, total)

    logger.info("Finished importing from %s", word_path)

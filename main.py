import os
import sys
import chardet
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QFileDialog, QTabWidget, QLineEdit, QMessageBox, QLabel
)

def detect_encoding(file_path, num_bytes=10000):
    """
    Определяет кодировку файла, читая первые num_bytes байт.
    Если обнаружено "MacCyrillic", заменяет её на "cp1251".
    """
    with open(file_path, 'rb') as f:
        rawdata = f.read(num_bytes)
    result = chardet.detect(rawdata)
    encoding = result.get('encoding')
    # Если chardet определил "MacCyrillic", считаем, что это cp1251
    if encoding and encoding.lower() == "maccyrillic":
        encoding = "cp1251"
    print(f"[DEBUG] Файл: {file_path}\n         Определена кодировка: {encoding}")
    return encoding

def read_lines_auto(file_path, default_encoding='utf-8', force_encoding=None):
    """
    Читает строки файла, используя либо принудительную кодировку,
    либо пытаясь определить её автоматически.
    При этом выводит отладочную информацию.
    """
    if force_encoding:
        encoding = force_encoding
        print(f"[DEBUG] Принудительная кодировка для {file_path}: {encoding}")
    else:
        encoding = detect_encoding(file_path) or default_encoding
        print(f"[DEBUG] Автоопределённая кодировка для {file_path}: {encoding}")
    try:
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            lines = [line.strip() for line in f.readlines()]
        print(f"[DEBUG] Считано {len(lines)} строк из файла: {file_path}")
        return lines
    except Exception as e:
        print(f"[ERROR] Ошибка чтения файла {file_path} с кодировкой {encoding}: {e}")
        with open(file_path, 'r', encoding=default_encoding, errors='replace') as f:
            lines = [line.strip() for line in f.readlines()]
        print(f"[DEBUG] Фолбэк: Считано {len(lines)} строк из файла: {file_path} с кодировкой {default_encoding}")
        return lines

def iter_block_items(parent):
    """
    Генератор, возвращающий последовательно параграфы и таблицы в документе.
    """
    if hasattr(parent, 'element'):
        parent_elm = parent.element.body
    else:
        parent_elm = parent
    for child in parent_elm:
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def export_to_word(eng_folder, rus_folder, output_path, rus_force_encoding=None):
    doc = Document()
    # Список файлов берём из папки с английскими файлами (имена должны совпадать)
    files = sorted([f for f in os.listdir(eng_folder) if f.lower().endswith('.txt')])
    if not files:
        raise Exception("Нет файлов .txt в папке с английскими файлами.")
    for filename in files:
        eng_file_path = os.path.join(eng_folder, filename)
        rus_file_path = os.path.join(rus_folder, filename)
        if not os.path.exists(rus_file_path):
            print(f"[WARNING] Для файла {filename} не найден соответствующий русский файл. Пропускаем.")
            continue

        # Чтение файлов:
        eng_lines = read_lines_auto(eng_file_path, default_encoding='utf-8')
        rus_lines = read_lines_auto(rus_file_path, default_encoding='cp1251', force_encoding=rus_force_encoding)

        # Добавляем заголовок с именем файла
        doc.add_paragraph(f'Файл: {filename}')

        # Создаём таблицу, где число строк – максимум из двух файлов
        num_rows = max(len(eng_lines), len(rus_lines))
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Table Grid'
        for i in range(num_rows):
            cell_eng = table.cell(i, 0)
            cell_rus = table.cell(i, 1)
            cell_eng.text = eng_lines[i] if i < len(eng_lines) else ''
            cell_rus.text = rus_lines[i] if i < len(rus_lines) else ''
        # Добавляем пустой параграф для разделения блоков
        doc.add_paragraph()
    doc.save(output_path)
    print(f"[DEBUG] Word-документ сохранён: {output_path}")

def import_from_word(word_path, eng_output_folder, rus_output_folder):
    doc = Document(word_path)
    current_filename = None
    file_data = {}  # Имя файла -> (англ. строки, рус. строки)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text.startswith("Файл:"):
                current_filename = text.replace("Файл:", "").strip()
                file_data[current_filename] = ([], [])
                print(f"[DEBUG] Начало секции для файла: {current_filename}")
        elif isinstance(block, Table) and current_filename:
            eng_lines, rus_lines = file_data[current_filename]
            for row in block.rows:
                eng_text = row.cells[0].text.strip()
                rus_text = row.cells[1].text.strip()
                eng_lines.append(eng_text)
                rus_lines.append(rus_text)
            current_filename = None
    os.makedirs(eng_output_folder, exist_ok=True)
    os.makedirs(rus_output_folder, exist_ok=True)
    for filename, (eng_lines, rus_lines) in file_data.items():
        filename_txt = filename if filename.lower().endswith('.txt') else filename + '.txt'
        eng_file_path = os.path.join(eng_output_folder, filename_txt)
        rus_file_path = os.path.join(rus_output_folder, filename_txt)
        with open(eng_file_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(eng_lines))
        with open(rus_file_path, 'w', encoding='cp1251') as f:
            f.write("\n".join(rus_lines))
        print(f"[DEBUG] Сохранены файлы: {eng_file_path} и {rus_file_path}")

# ----------------------- Интерфейс PyQt5 -----------------------

class ExportTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.eng_folder = ""
        self.rus_folder = ""
        self.rus_encoding = ""  # Если оставить пустым – автоопределение для каждого файла
        layout = QVBoxLayout()

        # Выбор папки с английскими файлами
        hbox_eng = QHBoxLayout()
        self.eng_folder_edit = QLineEdit()
        btn_select_eng = QPushButton("Выбрать папку с английскими файлами")
        btn_select_eng.clicked.connect(self.select_eng_folder)
        hbox_eng.addWidget(self.eng_folder_edit)
        hbox_eng.addWidget(btn_select_eng)
        layout.addLayout(hbox_eng)

        # Выбор папки с русскими файлами
        hbox_rus = QHBoxLayout()
        self.rus_folder_edit = QLineEdit()
        btn_select_rus = QPushButton("Выбрать папку с русскими файлами")
        btn_select_rus.clicked.connect(self.select_rus_folder)
        hbox_rus.addWidget(self.rus_folder_edit)
        hbox_rus.addWidget(btn_select_rus)
        layout.addLayout(hbox_rus)

        # Поле для указания кодировки русских файлов (оставьте пустым для автоопределения)
        hbox_encoding = QHBoxLayout()
        encoding_label = QLabel("Кодировка русских файлов (пусто = авто):")
        self.encoding_edit = QLineEdit(self.rus_encoding)
        hbox_encoding.addWidget(encoding_label)
        hbox_encoding.addWidget(self.encoding_edit)
        layout.addLayout(hbox_encoding)

        # Кнопка для создания Word-документа
        btn_export = QPushButton("Создать Word документ")
        btn_export.clicked.connect(self.do_export)
        layout.addWidget(btn_export)

        self.setLayout(layout)

    def select_eng_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку с английскими файлами")
        if folder:
            self.eng_folder = folder
            self.eng_folder_edit.setText(folder)

    def select_rus_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку с русскими файлами")
        if folder:
            self.rus_folder = folder
            self.rus_folder_edit.setText(folder)

    def do_export(self):
        if not self.eng_folder or not self.rus_folder:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите обе папки с файлами")
            return
        # Если поле пустое, будет использоваться автоопределение для каждого файла
        rus_enc = self.encoding_edit.text().strip() or None
        output_path, _ = QFileDialog.getSaveFileName(self,
                                                     "Сохранить Word документ",
                                                     "",
                                                     "Word файлы (*.docx)")
        if not output_path:
            return
        try:
            export_to_word(self.eng_folder, self.rus_folder, output_path, rus_force_encoding=rus_enc)
            QMessageBox.information(self, "Успех", "Word документ успешно создан!")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

class ImportTab(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.word_file = ""
        self.eng_output_folder = ""
        self.rus_output_folder = ""
        layout = QVBoxLayout()

        # Выбор Word-документа
        hbox_word = QHBoxLayout()
        self.word_file_edit = QLineEdit()
        btn_select_word = QPushButton("Выбрать Word документ")
        btn_select_word.clicked.connect(self.select_word_file)
        hbox_word.addWidget(self.word_file_edit)
        hbox_word.addWidget(btn_select_word)
        layout.addLayout(hbox_word)

        # Выбор папки для английских файлов
        hbox_eng = QHBoxLayout()
        self.eng_output_edit = QLineEdit()
        btn_select_eng_out = QPushButton("Папка для английских файлов")
        btn_select_eng_out.clicked.connect(self.select_eng_output)
        hbox_eng.addWidget(self.eng_output_edit)
        hbox_eng.addWidget(btn_select_eng_out)
        layout.addLayout(hbox_eng)

        # Выбор папки для русских файлов
        hbox_rus = QHBoxLayout()
        self.rus_output_edit = QLineEdit()
        btn_select_rus_out = QPushButton("Папка для русских файлов")
        btn_select_rus_out.clicked.connect(self.select_rus_output)
        hbox_rus.addWidget(self.rus_output_edit)
        hbox_rus.addWidget(btn_select_rus_out)
        layout.addLayout(hbox_rus)

        # Кнопка для импорта из Word
        btn_import = QPushButton("Разбить Word документ")
        btn_import.clicked.connect(self.do_import)
        layout.addWidget(btn_import)

        self.setLayout(layout)

    def select_word_file(self):
        file, _ = QFileDialog.getOpenFileName(self, "Выберите Word документ", "", "Word файлы (*.docx)")
        if file:
            self.word_file = file
            self.word_file_edit.setText(file)

    def select_eng_output(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для английских файлов")
        if folder:
            self.eng_output_folder = folder
            self.eng_output_edit.setText(folder)

    def select_rus_output(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку для русских файлов")
        if folder:
            self.rus_output_folder = folder
            self.rus_output_edit.setText(folder)

    def do_import(self):
        if not self.word_file or not self.eng_output_folder or not self.rus_output_folder:
            QMessageBox.warning(self, "Ошибка", "Сначала выберите Word документ и папки для сохранения файлов")
            return
        try:
            import_from_word(self.word_file, self.eng_output_folder, self.rus_output_folder)
            QMessageBox.information(self, "Успех", "Файлы успешно сохранены!")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Конвертер субтитров: TXT <-> Word")
        self.resize(600, 300)
        tabs = QTabWidget()
        tabs.addTab(ExportTab(), "Создать Word документ")
        tabs.addTab(ImportTab(), "Разбить Word документ")
        self.setCentralWidget(tabs)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
